<<<<<<< HEAD
﻿from __future__ import annotations
=======
from __future__ import annotations
>>>>>>> upstream/master

import io
import os
from pathlib import Path
<<<<<<< HEAD
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
=======
from typing import Iterable

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
>>>>>>> upstream/master
from pypdf import PdfReader

from PIL import Image
from docx import Document as DocxDocument

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None


TEXT_EXTENSIONS = {".txt", ".docx", ".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
THREED_EXTENSIONS = {".glb", ".gltf", ".obj"}

DEFAULT_CHUNK_SIZE = int(os.getenv("INGEST_CHUNK_SIZE", "2000"))
DEFAULT_CHUNK_OVERLAP = int(os.getenv("INGEST_CHUNK_OVERLAP", "250"))


<<<<<<< HEAD
def _get_splitter(chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def _split_text(text: str, filename: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[Document]:
    document = Document(page_content=text, metadata={"fuente": filename})
    return _get_splitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap).split_documents([document])
=======
def _split_text(text: str, filename: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    documentos = [Document(page_content=text, metadata={"fuente": filename})]
    return splitter.split_documents(documentos)
>>>>>>> upstream/master


def _texto_desde_txt(contenido: bytes) -> str:
    return contenido.decode("utf-8", errors="replace")


def _texto_desde_docx(contenido: bytes) -> str:
    documento = DocxDocument(io.BytesIO(contenido))
    parrafos = [p.text for p in documento.paragraphs if p.text.strip()]
    return "\n".join(parrafos)


<<<<<<< HEAD
def _texto_desde_pdf(contenido: bytes) -> List[Document]:
    reader = PdfReader(io.BytesIO(contenido))
    paginas: List[Document] = []
    for numero, pagina in enumerate(reader.pages, start=1):
        texto = pagina.extract_text() or ""
        if texto.strip():
            paginas.append(Document(page_content=texto.strip(), metadata={"fuente": f"pagina_{numero}"}))
    if not paginas:
        return []
    return _get_splitter().split_documents(paginas)


def _texto_desde_imagen(contenido: bytes, filename: str) -> List[Document]:
=======
def _texto_desde_pdf(contenido: bytes) -> list[Document]:
    reader = PdfReader(io.BytesIO(contenido))
    paginas = []
    for page_number, page in enumerate(reader.pages, start=1):
        texto = page.extract_text() or ""
        if texto.strip():
            paginas.append(Document(page_content=texto.strip(), metadata={"fuente": f"pagina_{page_number}"}))
    if not paginas:
        return []
    splitter = RecursiveCharacterTextSplitter(chunk_size=DEFAULT_CHUNK_SIZE, chunk_overlap=DEFAULT_CHUNK_OVERLAP)
    return splitter.split_documents(paginas)


def _texto_desde_imagen(contenido: bytes, filename: str) -> list[Document]:
>>>>>>> upstream/master
    try:
        imagen = Image.open(io.BytesIO(contenido))
        metadata = {k: str(v) for k, v in imagen.info.items() if v}
        ocr_text = ""
        if pytesseract is not None:
            try:
                ocr_text = pytesseract.image_to_string(imagen)
            except Exception:
                ocr_text = ""
        texto = (
            f"Imagen {filename} cargada. Metadata: {metadata or 'sin metadata disponible'}."
            f"\n\nTexto detectado: {ocr_text.strip() or 'No se detectó texto legible.'}"
        )
        return _split_text(texto, filename)
    except Exception:
        return [Document(page_content=f"Imagen {filename} no pudo procesarse correctamente.", metadata={"fuente": filename})]


<<<<<<< HEAD
def _texto_desde_modelo_3d(referencia: str, titulo: str | None = None, descripcion: str | None = None) -> List[Document]:
    cuerpo = [
        f"Modelo 3D: {titulo or referencia}",
        f"Referencia: {referencia}",
    ]
    if descripcion:
        cuerpo.append(f"Descripción: {descripcion}")
    cuerpo.append(
        "Este recurso debe ser tratado como un elemento interactivo de ingeniería de diseño, relacionado con materiales, planos o modelado 3D."
=======
def _texto_desde_modelo_3d(referencia: str, titulo: str | None = None, descripcion: str | None = None) -> list[Document]:
    cuerpo = [f"Modelo 3D: {titulo or referencia}", f"Referencia: {referencia}"]
    if descripcion:
        cuerpo.append(f"Descripción: {descripcion}")
    cuerpo.append(
        "Este recurso debe ser tratado como un elemento interactivo de ingeniería de diseño, relacionado con materiales, planos o modelado 3D."  # noqa: E501
>>>>>>> upstream/master
    )
    return _split_text("\n".join(cuerpo), referencia)


<<<<<<< HEAD
def cargar_documento(path: Path) -> List[Document]:
    extension = path.suffix.lower()
    contenido = path.read_bytes()

    if extension == ".pdf":
        return _texto_desde_pdf(contenido)
    if extension == ".txt":
        return _split_text(_texto_desde_txt(contenido), path.name)
    if extension == ".docx":
        return _split_text(_texto_desde_docx(contenido), path.name)
    if extension in IMAGE_EXTENSIONS:
        return _texto_desde_imagen(contenido, path.name)
=======
def cargar_documento(path: Path) -> list[Document]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        loader = PyPDFLoader(str(path))
        documentos = loader.load()
        return RecursiveCharacterTextSplitter(chunk_size=DEFAULT_CHUNK_SIZE, chunk_overlap=DEFAULT_CHUNK_OVERLAP).split_documents(documentos)
    if extension == ".txt":
        texto = _texto_desde_txt(path.read_bytes())
        return _split_text(texto, path.name)
    if extension == ".docx":
        texto = _texto_desde_docx(path.read_bytes())
        return _split_text(texto, path.name)
    if extension in IMAGE_EXTENSIONS:
        return _texto_desde_imagen(path.read_bytes(), path.name)
>>>>>>> upstream/master
    if extension in THREED_EXTENSIONS:
        return _texto_desde_modelo_3d(path.name, descripcion="Archivo 3D de referencia para ingeniería de diseño.")
    raise ValueError(f"Formato no soportado para la ingesta: {extension}")


<<<<<<< HEAD
def cargar_documento_desde_bytes(filename: str, content: bytes) -> List[Document]:
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return _texto_desde_pdf(content)
    if extension == ".txt":
        return _split_text(_texto_desde_txt(content), filename)
    if extension == ".docx":
        return _split_text(_texto_desde_docx(content), filename)
=======
def cargar_documento_desde_bytes(filename: str, content: bytes) -> list[Document]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _texto_desde_pdf(content)
    if extension == ".txt":
        texto = _texto_desde_txt(content)
        return _split_text(texto, filename)
    if extension == ".docx":
        texto = _texto_desde_docx(content)
        return _split_text(texto, filename)
>>>>>>> upstream/master
    if extension in IMAGE_EXTENSIONS:
        return _texto_desde_imagen(content, filename)
    if extension in THREED_EXTENSIONS:
        return _texto_desde_modelo_3d(filename, descripcion="Archivo 3D de referencia para ingeniería de diseño.")
    raise ValueError(f"Formato no soportado para la ingesta: {extension}")
